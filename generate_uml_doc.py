from docx import Document
from docx.shared import Pt


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def main():
    doc = Document()

    add_heading(doc, "Document UML final - Plateforme Freelancing", level=1)
    add_paragraph(doc, "Version: finale")

    add_heading(doc, "1. Acteurs", level=2)
    add_paragraph(doc, "Client")
    add_paragraph(doc, "Freelancer (heritage de Client)")
    add_paragraph(doc, "Admin (heritage de Freelancer et donc de Client)")
    add_paragraph(doc, "Service Paiement (acteur externe)")
    add_paragraph(doc, "Service Notifications (acteur externe)")

    add_heading(doc, "2. Diagramme de cas d'utilisation (Mermaid)", level=2)
    mermaid = """usecaseDiagram
title Plateforme Freelancing - Cas d'utilisation (Version finale)

actor Client as C
actor Freelancer as F
actor Admin as A
actor "Service Paiement" as PSP
actor "Service Notifications" as NS

F --|> C
A --|> F

rectangle "Systeme Plateforme Freelancing" {
  (Creer un compte) as UC_Register
  (Se connecter) as UC_Login
  (Verifier email) as UC_VerifyEmail
  (Reinitialiser mot de passe) as UC_ResetPwd
  (Recevoir notifications email/push) as UC_Notify
  (Creer / Mettre a jour profil) as UC_Profile
  (Definir taux horaire par defaut) as UC_DefaultRate
  (Soumettre mission) as UC_SubmitJob
  (Definir budget mission (fixe/horaire)) as UC_SetJobBudget
  (Mission en attente de moderation) as UC_PendingModeration
  (Valider mission) as UC_ApproveJob
  (Rejeter mission) as UC_RejectJob
  (Moderer mission) as UC_ModerateJob
  (Consulter missions) as UC_BrowseJobs
  (Postuler a une mission) as UC_Apply
  (Proposer montant / taux horaire) as UC_ProposePrice
  (Rechercher meilleure proposition) as UC_SearchBestProp
  (Valider proposition) as UC_AcceptProp
  (Rejeter proposition) as UC_RejectProp
  (Ignorer proposition) as UC_IgnoreProp
  (Prefinancer mission (escrow)) as UC_FundEscrow
  (Demarrer mission) as UC_StartMission
  (Cloturer mission) as UC_CloseMission
  (Transferer paiement au freelancer) as UC_ReleaseFunds
  (Ouvrir litige) as UC_OpenDispute
  (Traiter litige) as UC_HandleDispute
  (Rembourser client) as UC_RefundClient
  (Payer freelancer) as UC_PayFreelancer
  (Partager montant client/freelancer) as UC_SplitAmount
  (Prelever commission plateforme) as UC_TakeCommission
  (Laisser avis sur mission) as UC_ReviewMission
  (Noter client/freelancer) as UC_RateUser
  (Gerer utilisateurs) as UC_ManageUsers
}

C --> UC_Register
C --> UC_Login
C --> UC_VerifyEmail
C --> UC_ResetPwd
C --> UC_Profile
C --> UC_SubmitJob
C --> UC_SearchBestProp
C --> UC_AcceptProp
C --> UC_RejectProp
C --> UC_IgnoreProp
C --> UC_FundEscrow
C --> UC_CloseMission
C --> UC_OpenDispute
C --> UC_ReviewMission
C --> UC_RateUser
C --> UC_Notify

F --> UC_BrowseJobs
F --> UC_Apply
F --> UC_OpenDispute
F --> UC_ReviewMission
F --> UC_RateUser
F --> UC_DefaultRate
F --> UC_Notify

A --> UC_ModerateJob
A --> UC_ApproveJob
A --> UC_RejectJob
A --> UC_HandleDispute
A --> UC_ManageUsers

PSP --> UC_FundEscrow
PSP --> UC_ReleaseFunds
PSP --> UC_RefundClient
PSP --> UC_PayFreelancer
PSP --> UC_SplitAmount

NS --> UC_Notify

UC_SubmitJob ..> UC_SetJobBudget : <<include>>
UC_SubmitJob ..> UC_PendingModeration : <<include>>
UC_ModerateJob ..> UC_ApproveJob : <<extend>>
UC_ModerateJob ..> UC_RejectJob : <<extend>>
UC_Apply ..> UC_ProposePrice : <<include>>
UC_AcceptProp ..> UC_FundEscrow : <<include>>
UC_FundEscrow ..> UC_StartMission : <<include>>
UC_CloseMission ..> UC_ReleaseFunds : <<include>>
UC_ReleaseFunds ..> UC_TakeCommission : <<include>>
UC_HandleDispute ..> UC_RefundClient : <<extend>>
UC_HandleDispute ..> UC_PayFreelancer : <<extend>>
UC_HandleDispute ..> UC_SplitAmount : <<extend>>
UC_RefundClient ..> UC_TakeCommission : <<include>>
UC_PayFreelancer ..> UC_TakeCommission : <<include>>
UC_SplitAmount ..> UC_TakeCommission : <<include>>
UC_ReviewMission ..> UC_RateUser : <<include>>"""
    add_paragraph(doc, mermaid)

    add_heading(doc, "3. Regles metier", level=2)
    rules = [
        "Le client definit le budget mission (fixe ou horaire) a la soumission.",
        "Le freelancer propose son montant/taux lors de la candidature.",
        "La mission ne demarre qu'apres prefinancement (escrow).",
        "Sans litige: cloture mission => transfert automatique vers freelancer.",
        "En litige: l'admin decide remboursement client, paiement freelancer, ou partage.",
        "La plateforme preleve sa commission quelle que soit l'issue du litige.",
        "Les deux parties peuvent laisser un avis/notation en fin de mission.",
        "Notifications email/push pour les evenements majeurs (mission, proposition, litige, paiement).",
    ]
    for r in rules:
        add_paragraph(doc, f"- {r}")

    add_heading(doc, "4. Description detaillee des principaux cas d'utilisation", level=2)
    details = [
        ("Soumettre mission", "Le client cree une mission, precise le budget (fixe/horaire), puis la mission passe en attente de moderation."),
        ("Moderer mission", "L'admin verifie la mission et peut la valider ou la rejeter."),
        ("Postuler a mission", "Le freelancer consulte les missions valides et propose un montant ou un taux horaire."),
        ("Selectionner proposition", "Le client compare les propositions et peut valider, rejeter ou ignorer."),
        ("Prefinancer mission", "Avant demarrage, le client bloque les fonds en escrow via le service de paiement."),
        ("Cloturer mission", "A la fin, le client cloture la mission et le systeme transfere automatiquement les fonds au freelancer (moins commission)."),
        ("Ouvrir et traiter litige", "Client ou freelancer ouvre un litige; l'admin arbitre et applique la decision financiere."),
        ("Noter et laisser avis", "Chaque partie peut noter l'autre sur la mission terminee."),
    ]
    for title, desc in details:
        add_paragraph(doc, title + ":", bold=True)
        add_paragraph(doc, desc)

    output_path = "UML_Final_Plateforme_Freelancing.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
